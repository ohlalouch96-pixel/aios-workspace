from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Titel
title = doc.add_heading('Checklist — Oriënterend gesprek MHL Installatieservice', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Insightance | Juni 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

blokken = [
    ('Voor het gesprek', [
        'Afspraak ingepland en bevestigd met vader (datum + tijd)',
        'Google Meet of Zoom link aangemaakt en gedeeld',
        'Fathom actief en gekoppeld aan je Google/Zoom account',
        'Test Fathom even van tevoren met een kort proefgesprekje',
        'Intake-vragenlijst open op je scherm (het Word document)',
        'Notitieblok of tweede scherm bij de hand voor snelle aantekeningen',
        'Vader laten weten dat je het gesprek opneemt voor je aantekeningen',
    ]),
    ('Tijdens het gesprek', [
        'Fathom gestart en opname actief',
        'Vraag 1 t/m 5 doorlopen — rustig tempo, laat hem uitpraten',
        'Doorvragen waar je meer wilt begrijpen ("hoe bedoel je dat?", "hoe lang kost dat je?")',
        'Peilen of hij open staat voor verandering — niet aannemen, gewoon vragen',
        'Geen oplossingen aandragen — alleen luisteren en begrijpen',
    ]),
    ('Na het gesprek', [
        'python scripts/intel/collect_all.py draaien om samenvatting op te halen',
        'Samenvatting doorlezen en aanvullen waar nodig',
        'Belangrijkste pijnpunten noteren — dit wordt de basis voor stap 2 (analyse)',
    ]),
]

for blok_titel, items in blokken:
    doc.add_heading(blok_titel, level=2)
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('☐  ' + item)
    doc.add_paragraph('')

doc.save(r'c:\Users\oussa\Desktop\Checklist_MHL_Gesprek.docx')
print('Klaar')
