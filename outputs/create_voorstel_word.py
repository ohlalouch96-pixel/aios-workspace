from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Paginamarges instellen ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.2)
section.bottom_margin = Cm(2.2)

GOLD    = RGBColor(0xC9, 0xA8, 0x4C)
DARK    = RGBColor(0x0D, 0x0D, 0x10)
CREAM   = RGBColor(0xEA, 0xE5, 0xDC)
TEXT    = RGBColor(0x2D, 0x28, 0x25)
GRAY    = RGBColor(0x77, 0x77, 0x77)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
RED     = RGBColor(0xC0, 0x39, 0x2B)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def para(text='', bold=False, size=11, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size  = Pt(size)
        r.font.color.rgb = color
        r.font.name  = 'Calibri'
    return p

def label(text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = color
    r.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def heading(text, size=22, color=TEXT, space_before=4, space_after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(size)
    r.font.color.rgb = color
    r.font.name  = 'Calibri'
    return p

def intro(text, color=GRAY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)
    p.paragraph_format.left_indent  = Cm(0.5)
    # gouden lijn links via border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), 'C9A84C')
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.font.size  = Pt(11)
    r.font.color.rgb = color
    r.font.name  = 'Calibri'
    r.italic = True
    return p

def bullet(text, bold_prefix=None, color=TEXT, sub_color=GOLD, sub_text=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if bold_prefix:
        rb = p.add_run(bold_prefix + ' ')
        rb.bold = True
        rb.font.size = Pt(11)
        rb.font.color.rgb = TEXT
        rb.font.name = 'Calibri'
    r = p.add_run(text)
    r.font.size  = Pt(11)
    r.font.color.rgb = color
    r.font.name  = 'Calibri'
    if sub_text:
        rs = p.add_run('  ' + sub_text)
        rs.font.size  = Pt(9)
        rs.font.color.rgb = sub_color
        rs.font.name  = 'Calibri'
        rs.bold = True
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'C9A84C')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def block_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(11)
    r.font.color.rgb = TEXT
    r.font.name  = 'Calibri'
    return p

def page_break():
    doc.add_page_break()

def header_row(brand, num):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for cell in tbl.rows[0].cells:
        cell.width = Cm(8)
        for border_name in ['top','bottom','left','right']:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
            tcPr.append(tcBorders)
    c0, c1 = tbl.rows[0].cells
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(brand)
    r0.bold = True
    r0.font.size = Pt(8)
    r0.font.color.rgb = GOLD
    r0.font.name = 'Calibri'
    p0.paragraph_format.space_after = Pt(0)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run(num)
    r1.font.size = Pt(8)
    r1.font.color.rgb = GRAY
    r1.font.name = 'Calibri'
    p1.paragraph_format.space_after = Pt(0)
    divider()


# ════════════════════════════════
# PAGINA 1 — COVER
# ════════════════════════════════
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tbl.rows[0].cells[0]
set_cell_bg(cell, '0D0D10')
cell.width = Cm(16)

def cover_para(text, size=11, color=WHITE, bold=False, italic=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, sb=4, sa=4):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size  = Pt(size)
        r.font.color.rgb = color
        r.font.name  = 'Calibri'
    return p

# Verwijder eerste lege paragraaf in cell
cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

cover_para('INSIGHTANCE', size=10, color=GOLD, bold=True, sb=16, sa=2)
cover_para('AI IMPLEMENTATIEVOORSTEL #001', size=9, color=RGBColor(0xC9,0xA8,0x4C), bold=True, sb=2, sa=24)
cover_para('Slimmer werken.', size=32, color=CREAM, bold=True, sb=0, sa=2)
cover_para('Minder handmatig.', size=32, color=CREAM, bold=True, sb=0, sa=2)
cover_para('Meer tijd.', size=32, color=CREAM, bold=True, sb=0, sa=28)
cover_para('OPGESTELD VOOR', size=8, color=RGBColor(0x80,0x70,0x50), bold=True, sb=0, sa=4)
cover_para('De Vries Installatietechniek', size=20, color=CREAM, bold=True, sb=0, sa=4)
cover_para('Mark de Vries, Eigenaar', size=11, color=RGBColor(0x80,0x78,0x68), sb=0, sa=24)
cover_para('4 juni 2026  ·  Amsterdam                          Voorstel #001  ·  Vertrouwelijk',
           size=9, color=RGBColor(0x44,0x44,0x44), sb=0, sa=16)

page_break()

# ════════════════════════════════
# PAGINA 2 — UW SITUATIE
# ════════════════════════════════
header_row('INSIGHTANCE', '02 / 06')
label('01 — Uw situatie')
heading('Waar het nu knelt')
intro('Tijdens ons gesprek werd duidelijk dat een aanzienlijk deel van uw tijd opgaat aan taken die herhaalbaar en voorspelbaar zijn. Taken die u afleiden van het werk waar u goed in bent: installaties uitvoeren en klanten helpen.')

block_title('Terugkerende knelpunten')
bullet('Alle afspraken worden handmatig ingepland via de telefoon, ook tijdens lopende werkzaamheden', sub_text='±4 uur/week')
bullet('Offertes worden handmatig opgesteld en verstuurd, elke keer opnieuw', sub_text='±3 uur/week')
bullet('Openstaande offertes worden niet systematisch opgevolgd, waardoor klanten afhaken zonder reactie', sub_text='±2 uur/week')
bullet('Klanten worden na een klus nooit gevraagd om een review, waardoor het online profiel achterblijft', sub_text='0 reviews/maand')

divider()

block_title('Wat dit in de praktijk betekent')
bullet('U bent moeilijk bereikbaar voor nieuwe klanten terwijl u aan het werk bent, waardoor aanvragen mislopen')
bullet('Elke offerte kost onnodig veel tijd, terwijl de inhoud grotendeels hetzelfde is')
bullet('Potentiële klanten die geen reactie ontvangen gaan naar de concurrent')

page_break()

# ════════════════════════════════
# PAGINA 3 — WAT U KRIJGT
# ════════════════════════════════
header_row('INSIGHTANCE', '03 / 06')
label('02 — Wat u krijgt')
heading('Wat dit voor u betekent')
intro('Dit is wat er concreet verandert na de implementatie: in uw dagelijkse werk, voor uw klanten en voor uw bedrijfsresultaat.')

results = [
    ('Nooit meer een gemiste aanvraag',
     'Klanten plannen 24/7 zelf een afspraak in via uw website, ook \'s avonds en in het weekend. U hoeft de telefoon niet op te nemen voor elke boeking. Nieuwe aanvragen komen automatisch binnen terwijl u aan het werk bent.',
     '±4 uur bespaard per week'),
    ('Een offerte versturen kost u nog maar 5 minuten',
     'Op basis van uw intake wordt automatisch een professionele offerte aangemaakt. U controleert, past eventueel aan en verstuurt met één klik. Geen blanco document meer, geen herhalend typewerk.',
     '±3 uur bespaard per week'),
    ('Elke offerte wordt automatisch opgevolgd',
     'Heeft een klant na 3 dagen nog niet gereageerd? Dan stuurt het systeem automatisch een vriendelijke herinnering. U hoeft er niet meer aan te denken. Geen verloren klanten meer door stilte.',
     'Hogere offerteconversie zonder extra inspanning'),
    ('Meer Google-reviews zonder er iets voor te doen',
     'Na elke afgeronde klus ontvangt de klant automatisch een reviewverzoek via WhatsApp of e-mail. Tevreden klanten klikken in 30 seconden op de link. Uw online profiel groeit vanzelf.',
     'Gemiddeld 3 tot 5 nieuwe reviews per maand'),
]

for i, (title, desc, metric) in enumerate(results, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    rn = p.add_run(f'{i}.  ')
    rn.bold = True
    rn.font.size = Pt(13)
    rn.font.color.rgb = GOLD
    rn.font.name = 'Calibri'
    rt = p.add_run(title)
    rt.bold = True
    rt.font.size = Pt(13)
    rt.font.color.rgb = TEXT
    rt.font.name = 'Calibri'
    para(desc, color=GRAY, size=11, space_before=2, space_after=2)
    pm = doc.add_paragraph()
    pm.paragraph_format.space_before = Pt(2)
    pm.paragraph_format.space_after  = Pt(6)
    rm = pm.add_run(f'→  {metric}')
    rm.bold = True
    rm.font.size = Pt(10)
    rm.font.color.rgb = GOLD
    rm.font.name = 'Calibri'

page_break()

# ════════════════════════════════
# PAGINA 4 — ONZE AANPAK
# ════════════════════════════════
header_row('INSIGHTANCE', '04 / 06')
label('03 — Onze aanpak')
heading('Hoe we dit aanpakken')
intro('Van start tot oplevering. U wordt bij elke stap betrokken en weet altijd waar we staan.')

steps = [
    ('1', 'Intake en analyse', '2 dagen',
     'We brengen uw huidige processen gedetailleerd in kaart. Wat zijn de stappen, waar zit de meeste tijdsinvestering, wat zijn de randvoorwaarden.'),
    ('2', 'Ontwerp en afstemming', '2 dagen',
     'We presenteren het ontwerp van het systeem. U geeft feedback en we stemmen af totdat alles klopt.'),
    ('3', 'Bouwen en testen', '2 weken',
     'Het systeem wordt gebouwd en uitgebreid getest. We simuleren echte situaties (afspraken inplannen, offertes versturen, reviews opvragen) zodat alles werkt zoals verwacht vóór de oplevering.'),
    ('4', 'Oplevering en instructie', '1 dag',
     'We leveren op, lopen samen door het systeem en zorgen dat u precies weet hoe u het gebruikt. Vragen achteraf worden altijd beantwoord.'),
]

for num, title, duration, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    rn = p.add_run(f'{num}.  ')
    rn.bold = True
    rn.font.size = Pt(14)
    rn.font.color.rgb = GOLD
    rn.font.name = 'Calibri'
    rt = p.add_run(title)
    rt.bold = True
    rt.font.size = Pt(13)
    rt.font.color.rgb = TEXT
    rt.font.name = 'Calibri'
    rd = p.add_run(f'  ({duration})')
    rd.font.size = Pt(10)
    rd.font.color.rgb = GOLD
    rd.font.name = 'Calibri'
    para(desc, color=GRAY, size=11, space_before=2, space_after=4)

page_break()

# ════════════════════════════════
# PAGINA 5 — INVESTERING
# ════════════════════════════════
header_row('INSIGHTANCE', '05 / 06')
label('04 — Investering')
heading('Wat het kost')
intro('Een transparant overzicht van de investering. Geen verrassingen achteraf.')

block_title('AI Implementatie op maat')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(2)
rp = p.add_run('€1.250')
rp.bold = True
rp.font.size = Pt(36)
rp.font.color.rgb = TEXT
rp.font.name = 'Calibri'
rs = p.add_run('  eenmalig')
rs.font.size = Pt(13)
rs.font.color.rgb = GRAY
rs.font.name = 'Calibri'
para('Exclusief BTW', color=GRAY, size=10, space_before=0, space_after=10)

bullet('Volledig werkend automatiseringssysteem (afspraken, offertes, reviews)')
bullet('Persoonlijke instructiesessie na oplevering')
bullet('30 dagen nazorg waarbij vragen altijd worden beantwoord')
bullet('Alle licentiekosten inbegrepen gedurende de eerste maand')

divider()

block_title('Betalingsvoorwaarden')
bullet('Na het startgesprek ontvangt u een officiële offerte. Zodra de aanbetaling van 50% is ontvangen, beginnen wij met bouwen.', bold_prefix='50% bij start.')
bullet('Het resterende bedrag is verschuldigd op het moment van oplevering van het systeem.', bold_prefix='50% bij oplevering.')

para('Dit voorstel is geldig tot 18 juni 2026. Na akkoord plannen we binnen 24 uur een startgesprek in.',
     color=GRAY, size=10, space_before=10, space_after=0)

page_break()

# ════════════════════════════════
# PAGINA 6 — VOLGENDE STAP
# ════════════════════════════════
header_row('INSIGHTANCE', '06 / 06')
label('05 — Volgende stap')
heading('Zo gaan we van start')
intro('Vier stappen scheiden u van een werkend systeem. Volg ze op volgorde, zo verliezen we geen tijd.')

action_steps = [
    ('1', 'Onderteken dit voorstel en stuur het retour',
     'Zet uw handtekening onderaan dit document en mail het terug naar oussama@insightance.ai. Vermeld in de onderwerpregel: Akkoord voorstel #001.'),
    ('2', 'Stuur minimaal 3 data\'s mee voor het startgesprek',
     'Voeg bij uw mail minimaal 3 momenten toe waarop u beschikbaar bent voor een startgesprek van 30 minuten. Wij bevestigen binnen 24 uur. Zo plannen we direct in. Geen heen-en-weer mailen, geen vertraging.'),
    ('3', 'Startgesprek: afstemming en officiële offerte',
     'Tijdens dit gesprek lopen we alles samen door zodat we zeker weten dat we op één lijn zitten. Na afloop ontvangt u de officiële offerte. Zodra de aanbetaling van 50% is voldaan, starten wij direct met bouwen.'),
    ('4', 'Oplevering en slotbetaling',
     'U ontvangt het systeem kant-en-klaar binnen 3 weken na de startdatum. Bij oplevering wordt het resterende bedrag van 50% gefactureerd.'),
]

for num, title, desc in action_steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    rn = p.add_run(f'{num}.  ')
    rn.bold = True
    rn.font.size = Pt(14)
    rn.font.color.rgb = GOLD
    rn.font.name = 'Calibri'
    rt = p.add_run(title)
    rt.bold = True
    rt.font.size = Pt(12)
    rt.font.color.rgb = TEXT
    rt.font.name = 'Calibri'
    para(desc, color=GRAY, size=11, space_before=2, space_after=4)

divider()

# Vragen box
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(2)
p.paragraph_format.left_indent  = Cm(0.5)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
left = OxmlElement('w:left')
left.set(qn('w:val'),   'single')
left.set(qn('w:sz'),    '16')
left.set(qn('w:space'), '4')
left.set(qn('w:color'), 'C9A84C')
pBdr.append(left)
pPr.append(pBdr)
rv = p.add_run('Vragen naar aanleiding van dit voorstel? ')
rv.bold = True
rv.font.size = Pt(11)
rv.font.color.rgb = TEXT
rv.font.name = 'Calibri'
rv2 = p.add_run('Stuur een mail naar oussama@insightance.ai met als onderwerp: Vraag voorstel #001. Wij reageren binnen één werkdag.')
rv2.font.size = Pt(11)
rv2.font.color.rgb = GRAY
rv2.font.name = 'Calibri'

# Handtekeningen
doc.add_paragraph().paragraph_format.space_after = Pt(20)

tbl2 = doc.add_table(rows=1, cols=2)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in tbl2.rows[0].cells:
    for border_name in ['top','bottom','left','right']:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        tcPr.append(tcBorders)

c0, c1 = tbl2.rows[0].cells

def sig_cell(cell, label_text, pre_sign, name, role, num=None):
    pl = cell.add_paragraph()
    rl = pl.add_run(label_text.upper())
    rl.font.size = Pt(7)
    rl.font.color.rgb = GRAY
    rl.font.name = 'Calibri'
    pl.paragraph_format.space_after = Pt(20)

    if pre_sign:
        pp = cell.add_paragraph()
        rp = pp.add_run(pre_sign)
        rp.italic = True
        rp.font.size = Pt(18)
        rp.font.color.rgb = TEXT
        rp.font.name = 'Georgia'
        pp.paragraph_format.space_after = Pt(2)

    pline = cell.add_paragraph()
    pPr2 = pline._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'E0DAD2')
    pBdr2.append(bot)
    pPr2.append(pBdr2)
    pline.paragraph_format.space_after = Pt(4)

    pn = cell.add_paragraph()
    rn = pn.add_run(name)
    rn.font.size = Pt(12)
    rn.font.color.rgb = TEXT
    rn.font.name = 'Calibri'
    rn.bold = True
    pn.paragraph_format.space_after = Pt(2)

    pr = cell.add_paragraph()
    rr = pr.add_run(role)
    rr.font.size = Pt(10)
    rr.font.color.rgb = GRAY
    rr.font.name = 'Calibri'
    pr.paragraph_format.space_after = Pt(2)

    if num:
        pnum = cell.add_paragraph()
        rnum = pnum.add_run(num)
        rnum.font.size = Pt(9)
        rnum.font.color.rgb = GOLD
        rnum.font.name = 'Calibri'

    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

sig_cell(c0, 'Akkoord opdrachtgever', None,
         'Mark de Vries', 'Eigenaar · De Vries Installatietechniek')
sig_cell(c1, 'Namens Insightance', 'O. Hlalouch',
         'Oussama Hlalouch', 'Oprichter · Insightance', 'Voorstel #001')

output = r'c:\Users\oussa\Desktop\Insightance_Voorstel_001.docx'
doc.save(output)
print(f'Opgeslagen: {output}')
