from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Stijlen instellen
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

# Titel
title = doc.add_heading('MHL Installatieservice — Oriënterende Analyse & Intake', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Opgesteld door: Insightance | Datum: juni 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# DEEL 1: ANALYSE
heading1('Deel 1 — Bedrijfsanalyse')

heading2('Bedrijfsprofiel')
body('MHL Installatieservice is een loodgietersbedrijf gevestigd in Amsterdam met meer dan 30 jaar ervaring. Het bedrijf bedient particuliere en zakelijke klanten in Amsterdam, Amstelveen, Zaandam, Haarlem en omstreken.')

heading2('Diensten')
for d in [
    'Loodgieterswerkzaamheden',
    'CV-ketelonderhoud en reparatie',
    'Ontstoppingsdiensten',
    'Lekkagereparatie',
    'Badkamer- en toiletrenovaties',
    'Dak- en zinkwerk',
    'Verwarming',
    '24/7 calamiteitendiensten',
]:
    bullet(d)

heading2('Waargenomen pijnpunten')
body('Op basis van de website zijn de volgende knelpunten geïdentificeerd:')
for p in [
    'Volledig handmatig aanvraagproces — klanten bellen direct voor offerte en planning',
    '24/7 calamiteitendienst betekent dat de eigenaar ook \'s nachts bereikbaar moet zijn',
    'Geen online aanvraagformulier of zelfplanningssysteem aanwezig',
    'Geen reviews of testimonials zichtbaar op de website — gemiste kans voor vertrouwen',
    'Website niet actief bijgehouden (placeholder-tekst nog aanwezig)',
]:
    bullet(p)

heading2('Potentiële automatiseringskansen')
body('De volgende processen komen in aanmerking voor automatisering:')
kansen = [
    ('Afsprakenplanning', 'Klanten zelf een tijdslot laten kiezen via een online boekingssysteem'),
    ('Offerte-aanvragen', 'Via een digitaal formulier in plaats van telefoon, automatisch doorgestuurd'),
    ('Calamiteiten triage', 'Een chatbot die \'s nachts eerste vragen stelt en urgentie bepaalt'),
    ('Reviews verzamelen', 'Automatisch een reviewverzoek sturen na afronding van een klus'),
    ('Factuuropvolging', 'Automatische herinnering als een offerte of factuur niet beantwoord wordt'),
]
for titel, beschrijving in kansen:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{titel}: ')
    run.bold = True
    p.add_run(beschrijving)

doc.add_paragraph('')

# DEEL 2: INTAKE-VRAGENLIJST
heading1('Deel 2 — Intake-vragenlijst')
body('Gebruik deze vragen tijdens het oriënterende gesprek. Doel is luisteren en begrijpen — nog geen oplossingen aandragen.')
doc.add_paragraph('')

blokken = [
    ('Blok 1 — Het bedrijf begrijpen', [
        'Hoeveel klussen doe je gemiddeld per week?',
        'Werk je alleen of heb je mensen in dienst?',
        'Wat zijn je drukste maanden en rustperiodes?',
    ]),
    ('Blok 2 — Het aanvraagproces', [
        'Hoe komen klanten nu bij jou terecht? (telefoon, via via, website?)',
        'Hoe plan je een afspraak in — doe je dat zelf, via de telefoon?',
        'Hoeveel tijd kost je dat gemiddeld per aanvraag?',
        'Heb je weleens aanvragen gemist omdat je in gesprek was of bezig was?',
    ]),
    ('Blok 3 — Offertes & administratie', [
        'Hoe maak je een offerte? (Word, app, uit je hoofd?)',
        'Hoelang duurt het van aanvraag tot verstuurde offerte?',
        'Wat doe je als een klant niet reageert op een offerte?',
        'Hoe stuur je facturen, en volg je openstaande facturen op?',
    ]),
    ('Blok 4 — Na de klus', [
        'Vraag je klanten om een review na een klus?',
        'Heb je vaste klanten die terugkomen — hoe houd je dat bij?',
    ]),
    ('Blok 5 — Pijn & prioriteit', [
        'Wat kost je op dit moment de meeste tijd buiten het échte werk?',
        'Wat zou je het liefst kwijt willen — wat ergert je het meest?',
        'Heb je al eens nagedacht over software of apps om dit te verbeteren?',
    ]),
]

for blok_titel, vragen in blokken:
    heading2(blok_titel)
    for i, vraag in enumerate(vragen, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(vraag)
    doc.add_paragraph('')

# Footer notitie
doc.add_paragraph('─' * 60)
body('Opmerking: Dit document is opgesteld op basis van openbare website-informatie. De intake-vragenlijst is bedoeld als leidraad — volg het gesprek en stel door waar nodig.')

output_path = r'c:\Users\oussa\Desktop\MHL_Analyse_en_Intake.docx'
doc.save(output_path)
print(f'Document opgeslagen: {output_path}')
